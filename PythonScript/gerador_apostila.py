import re
import io
import os
import json
import time # Para medir o tempo
import requests # pip install requests
from docx import Document # pip install python-docx
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import traceback
import concurrent.futures # Para ThreadPoolExecutor

# --- Configurações ---
INPUT_FILENAME = 'apostila_input.txt'
OUTPUT_FILENAME = 'apostila_gerada_paralela.docx' # Novo nome
RAYSO_API_URL = 'http://localhost:3000/api' # URL da API (via POST JSON)

# Parâmetros padrão para a API Rayso (width/padding como STRINGS)
RAYSO_DEFAULT_PARAMS = {
    "theme": "mintlify",   # String
    "darkMode": True,      # Boolean
    "padding": "16",       # String
    "background": True,   # Boolean - VOLTEI PARA TRUE CONFORME SEU ÚLTIMO CÓDIGO ENVIADO
    "width": "720",        # String
}

# Largura da imagem NO DOCUMENTO DOCX (em polegadas)
IMAGE_WIDTH_INCHES = 6.0

# Número máximo de chamadas paralelas à API Rayso
# Ajuste conforme a capacidade da sua máquina/API local (comece com baixo valor)
MAX_API_WORKERS = 1

# --- Fim das Configurações ---

CODE_BLOCK_REGEX = re.compile(r'\[CODE_BLOCK\s+(?:lang="(?P<lang>.*?)"\s*)?(?:title="(?P<title>.*?)"\s*)?(?:lang="(?P<lang2>.*?)"\s*)?(?:title="(?P<title2>.*?)"\s*)?\]', re.IGNORECASE)
FORMATTING_REGEX = re.compile(r'(\*\*.*?\*\*|\*.*?\*|\`.*?\`)')
INLINE_CODE_COLOR = RGBColor(0x00, 0x80, 0x00)
HEADING_COLOR = RGBColor(0x00, 0x00, 0x00)

# --- Função call_rayso_api (adaptada para logar ID) ---
def call_rayso_api(code, lang, title, block_id):
    """Chama a API Rayso local (via POST JSON) para gerar a imagem."""
    payload = {
        **RAYSO_DEFAULT_PARAMS,
        "code": code,
        "language": str(lang) if lang else "auto",
        "title": str(title) if title else f"snippet_{block_id}" # Usa ID se não houver título
    }
    # Garantir tipos
    if 'padding' in payload: payload['padding'] = str(payload['padding'])
    if 'width' in payload: payload['width'] = str(payload['width'])
    if 'darkMode' in payload: payload['darkMode'] = bool(payload['darkMode'])
    if 'background' in payload: payload['background'] = bool(payload['background'])

    print(f"  [Thread-{block_id}] Enviando POST para Rayso: title='{payload['title']}'...")
    response = None
    start_time = time.time()
    try:
        response = requests.post(RAYSO_API_URL, json=payload, timeout=60) # Timeout maior?
        response.raise_for_status()
        content_type = response.headers.get('content-type', '')
        if 'image' in content_type:
            end_time = time.time()
            print(f"  [Thread-{block_id}] Imagem '{payload['title']}' recebida com sucesso ({end_time - start_time:.2f}s).")
            return response.content # Sucesso -> retorna bytes da imagem
        else:
            print(f"  !! [Thread-{block_id}] ERRO: Resposta não é imagem para '{payload['title']}'. CT: {content_type}")
            # Logar resposta de erro pode ser útil
            try: print(f"     Resposta JSON: {response.json()}")
            except: print(f"     Resposta (não JSON): {response.text[:100]}...")
            return None # Falha -> retorna None
    except requests.exceptions.Timeout:
        print(f"  !! [Thread-{block_id}] ERRO: Timeout (POST) para '{payload['title']}'.")
    except requests.exceptions.RequestException as e:
        print(f"  !! [Thread-{block_id}] ERRO (POST) para '{payload['title']}': {e}")
    return None # Falha -> retorna None
# --- Fim da Função API ---

# --- Funções de formatação e adição ao DOCX (mantidas) ---
def add_formatted_run(paragraph, text_part):
    if not text_part: return
    content_bold = text_part[2:-2] if len(text_part) > 4 else ''
    content_italic = text_part[1:-1] if len(text_part) > 2 else ''
    content_code = text_part[1:-1] if len(text_part) > 2 else ''
    if text_part.startswith('**') and text_part.endswith('**'):
        run = paragraph.add_run(content_bold if content_bold else text_part)
        if content_bold: run.bold = True
    elif text_part.startswith('*') and text_part.endswith('*'):
        run = paragraph.add_run(content_italic if content_italic else text_part)
        if content_italic: run.italic = True
    elif text_part.startswith('`') and text_part.endswith('`'):
        run = paragraph.add_run(content_code if content_code else text_part)
        if content_code: run.font.color.rgb = INLINE_CODE_COLOR
    else:
        paragraph.add_run(text_part)

def add_formatted_paragraph(doc_or_para, text):
    if isinstance(doc_or_para, DocxDocument):
        paragraph = doc_or_para.add_paragraph()
    elif isinstance(doc_or_para, Paragraph):
        paragraph = doc_or_para
    else:
        print(f" !! ERRO INTERNO: add_formatted_paragraph tipo inesperado: {type(doc_or_para)}")
        return
    parts = FORMATTING_REGEX.split(text)
    for part in parts:
        add_formatted_run(paragraph, part)

def add_colored_heading(document, text, level):
    heading_paragraph = document.add_heading(text, level=level)
    for run in heading_paragraph.runs:
        run.font.color.rgb = HEADING_COLOR
    return heading_paragraph
# --- Fim das funções auxiliares DOCX ---

# --- Função Principal Refatorada ---
def parse_and_generate_docx_parallel():
    """Lê o arquivo, processa blocos de código em paralelo e gera o DOCX."""
    if not os.path.exists(INPUT_FILENAME):
        print(f" !! ERRO: Arquivo de entrada '{INPUT_FILENAME}' não encontrado.")
        return

    print(f"Iniciando geração paralela do documento '{OUTPUT_FILENAME}'...")
    start_total_time = time.time()

    # --- FASE 1: Parsear o arquivo e estruturar o conteúdo ---
    print("[FASE 1] Lendo e parseando o arquivo de entrada...")
    content_structure = [] # Lista para guardar a estrutura: {'type': '...', 'data': ...}
    code_blocks_to_process = [] # Lista para tarefas da API: {'id': ..., 'lang': ..., 'title': ..., 'code': ...}
    code_block_counter = 0

    in_code_block = False
    current_code = []
    code_lang = None
    code_title = None
    line_number = 0

    try:
        with open(INPUT_FILENAME, 'r', encoding='utf-8') as f:
            for line_number, line in enumerate(f, 1):
                original_line = line
                stripped_line = line.strip()

                match = CODE_BLOCK_REGEX.match(line)
                if match and not in_code_block:
                    code_lang = match.group('lang') or match.group('lang2')
                    code_title = match.group('title') or match.group('title2')
                    in_code_block = True
                    current_code = []
                    continue # Próxima linha é código

                if '[/CODE_BLOCK]' in line and in_code_block:
                    in_code_block = False
                    code_content = "\n".join(current_code).strip()
                    code_block_counter += 1
                    block_data = {
                        'id': code_block_counter,
                        'lang': code_lang,
                        'title': code_title,
                        'code': code_content
                    }
                    if code_content:
                        code_blocks_to_process.append(block_data)
                        content_structure.append({'type': 'code_block_placeholder', 'id': code_block_counter, 'title': code_title})
                    else:
                         print(f" !! AVISO (Linha {line_number}): Bloco de código '{code_title}' vazio ignorado.")
                    code_lang, code_title = None, None # Reset
                    continue

                if in_code_block:
                    current_code.append(original_line.rstrip('\n'))
                    continue

                # Fora de bloco de código - Adicionar à estrutura
                if stripped_line.startswith('### '):
                    content_structure.append({'type': 'heading', 'level': 3, 'text': stripped_line[4:].strip()})
                elif stripped_line.startswith('## '):
                    content_structure.append({'type': 'heading', 'level': 2, 'text': stripped_line[3:].strip()})
                elif stripped_line.startswith('# '):
                    content_structure.append({'type': 'heading', 'level': 1, 'text': stripped_line[2:].strip()})
                elif stripped_line.startswith(('-', '*')) and len(stripped_line) > 1 and stripped_line[1] == ' ':
                    content_structure.append({'type': 'list_item', 'text': stripped_line[2:].strip()})
                elif stripped_line:
                    content_structure.append({'type': 'paragraph', 'text': stripped_line})

    except FileNotFoundError:
         print(f" !! ERRO: Arquivo de entrada '{INPUT_FILENAME}' não encontrado durante leitura.")
         return
    except Exception as e:
        print(f" !! ERRO inesperado durante o parsing (linha ~{line_number}): {e}")
        traceback.print_exc()
        return

    print(f"[FASE 1] Parsing concluído. {len(content_structure)} elementos estruturais encontrados.")
    print(f"[FASE 1] {len(code_blocks_to_process)} blocos de código para gerar imagens.")

    # --- FASE 2: Gerar imagens em paralelo ---
    print(f"\n[FASE 2] Iniciando geração de imagens em paralelo (max_workers={MAX_API_WORKERS})...")
    image_results = {} # Dicionário para guardar resultados: {block_id: image_bytes or None}
    start_api_time = time.time()

    # Usando ThreadPoolExecutor para paralelizar
    with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_API_WORKERS) as executor:
        # Mapeia cada bloco de código para uma chamada futura à API
        future_to_block = {
            executor.submit(call_rayso_api, block['code'], block['lang'], block['title'], block['id']): block
            for block in code_blocks_to_process
        }

        # Processa os resultados conforme eles ficam prontos
        for future in concurrent.futures.as_completed(future_to_block):
            block_data = future_to_block[future]
            block_id = block_data['id']
            try:
                # Pega o resultado da thread (bytes da imagem ou None)
                image_bytes = future.result()
                image_results[block_id] = image_bytes # Armazena o resultado
            except Exception as exc:
                print(f" !! [Thread-{block_id}] ERRO INESPERADO ao processar futuro para '{block_data.get('title', block_id)}': {exc}")
                image_results[block_id] = None # Marca como falha

    end_api_time = time.time()
    print(f"[FASE 2] Geração de imagens concluída em {end_api_time - start_api_time:.2f} segundos.")
    # Contar sucessos e falhas
    success_count = sum(1 for res in image_results.values() if res is not None)
    fail_count = len(code_blocks_to_process) - success_count
    print(f"[FASE 2] {success_count} imagens geradas com sucesso, {fail_count} falhas.")

    # --- FASE 3: Montar o Documento DOCX ---
    print("\n[FASE 3] Montando o documento DOCX...")
    document = Document()
    start_docx_time = time.time()

    try:
        for item in content_structure:
            item_type = item.get('type')
            if item_type == 'heading':
                add_colored_heading(document, item.get('text', ''), item.get('level', 1))
            elif item_type == 'paragraph':
                add_formatted_paragraph(document, item.get('text', ''))
            elif item_type == 'list_item':
                p_list = document.add_paragraph(style='List Bullet')
                add_formatted_paragraph(p_list, item.get('text', ''))
            elif item_type == 'code_block_placeholder':
                block_id = item.get('id')
                block_title = item.get('title', f'snippet_{block_id}')
                image_bytes = image_results.get(block_id) # Pega resultado da Fase 2

                if image_bytes:
                    try:
                        image_stream = io.BytesIO(image_bytes)
                        document.add_picture(image_stream, width=Inches(IMAGE_WIDTH_INCHES))
                        document.add_paragraph() # Espaçamento
                    except Exception as e:
                        print(f" !! ERRO (DOCX): Ao inserir imagem ID {block_id} ('{block_title}') no DOCX: {e}")
                        p_err = document.add_paragraph()
                        add_formatted_run(p_err, f"*[Erro ao inserir imagem: {block_title}]*")
                else:
                    # Falha na geração da imagem (já logado na Fase 2)
                    p_fail = document.add_paragraph()
                    add_formatted_run(p_fail, f"*[Falha ao gerar imagem: {block_title}]*")

    except Exception as e:
        print(f" !! ERRO inesperado durante a montagem do DOCX: {e}")
        traceback.print_exc()
        return # Aborta se der erro na montagem

    end_docx_time = time.time()
    print(f"[FASE 3] Montagem do DOCX concluída em {end_docx_time - start_docx_time:.2f} segundos.")

    # Salva o documento final
    try:
        document.save(OUTPUT_FILENAME)
        end_total_time = time.time()
        print(f"\nDocumento '{OUTPUT_FILENAME}' gerado com sucesso!")
        print(f"Tempo total de execução: {end_total_time - start_total_time:.2f} segundos.")
    except Exception as e:
        print(f" !! ERRO ao salvar o documento '{OUTPUT_FILENAME}': {e}")

# --- Execução Principal ---
if __name__ == "__main__":
    parse_and_generate_docx_parallel()